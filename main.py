import random

from PyQt5 import QtCore, QtWidgets
import mysql.connector
from docx import Document

db_login = ""
db_pass = ""
db_host = ""
special_id = ''

first_column_dog = ['100001', '100002']
second_column_dog = ['26-12-2018', '23-01-2019']
third_column_dog = ['Пушкина 8', 'Колотушкина 9']

first_column_chek = ['110001', '110002']
second_column_chek = ['26-12-2018', '23-01-2019']

first_column_ord = ['120001', '120002']
second_column_ord = ['26-12-2018', '23-01-2019']
third_column_ord = ['2', '3']
four_column_ord = ['10000', '8000']


class Ui_MainWindow(object):
    def setupadminisUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(741, 380)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 0, 0, 1, 1)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 0, 1, 1, 4)
        self.pushButton_4 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_4.setObjectName("pushButton_4")
        self.gridLayout.addWidget(self.pushButton_4, 1, 7, 1, 1)
        self.pushButton_9 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_9.setObjectName("pushButton_9")
        self.gridLayout.addWidget(self.pushButton_9, 1, 3, 1, 3)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 1, 0, 1, 3)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 2, 0, 1, 2)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 2, 2, 1, 1)
        self.pushButton_6 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_6.setObjectName("pushButton_6")
        self.gridLayout.addWidget(self.pushButton_6, 1, 8, 1, 1)
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setObjectName("label_7")
        self.gridLayout.addWidget(self.label_7, 2, 5, 1, 1)
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 721, 253))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.gridLayout.addWidget(self.scrollArea, 3, 0, 1, 9)
        self.pushButton_3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_3.setObjectName("pushButton_3")
        self.gridLayout.addWidget(self.pushButton_3, 1, 6, 1, 1)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 741, 20))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateadminisUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateadminisUi(self, adminis):
        _translate = QtCore.QCoreApplication.translate
        adminis.setWindowTitle(_translate("adminis", "MainWindow"))
        self.label.setText(_translate("adminis", "Вы авторизованы как:"))
        self.label_2.setText(_translate("adminis", "администратор"))
        self.pushButton_9.setText(_translate("adminis", "Сформировать отчет"))
        self.label_3.setText(_translate("adminis",
                                        "<html><head/><body><p><span style=\" font-size:12pt;\">Список составленных договоров:</span></p></body></html>"))
        self.pushButton_4.setText(_translate("adminis", "Закрыть"))
        self.pushButton_6.setText(_translate("adminis", "Добавить "))
        self.label_5.setText(_translate("adminis", "Дата заключения"))
        self.label_4.setText(_translate("adminis", "Номер договора"))
        self.label_7.setText(_translate("adminis", "Адрес склада"))
        self.pushButton_3.setText(_translate("adminis", "Список сотрудников"))

        self.pushButton_3.clicked.connect(self.setupworkersUi)
        self.pushButton_6.clicked.connect(lambda: add())
        self.pushButton_9.clicked.connect(self.openotchet)

        def add():
             first_column_dog.append('100003')
             second_column_dog.append('24-01-2019')
             third_column_dog.append('Песочная 14')
             self.setupadddogovorUi()

        i = 0
        for item in first_column_dog:
            line_item = QtWidgets.QLabel(str(item))
            self.gridLayout_2.addWidget(line_item, i, 1, 1, 1)
            i += 1
        i = 0
        for item in second_column_dog:
            line_item = QtWidgets.QLineEdit(str(item))
            self.gridLayout_2.addWidget(line_item, i, 2, 1, 1)
            i += 1
        i = 0
        for item in third_column_dog:
            line_item = QtWidgets.QLineEdit(str(item))
            self.gridLayout_2.addWidget(line_item, i, 3, 1, 1)
            line_item = QtWidgets.QPushButton("Удалить")
            line_item.clicked.connect(lambda state, row = i: delete(row))
            self.gridLayout_2.addWidget(line_item, i, 4, 1, 1)
            i += 1

        def delete(row):
            first_column_dog.pop(row)
            second_column_dog.pop(row)
            third_column_dog.pop(row)
            self.setupadminisUi()

    def setupadddogovorUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(313, 430)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setObjectName("lineEdit")
        self.verticalLayout.addWidget(self.lineEdit)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.verticalLayout.addWidget(self.lineEdit_2)
        self.lineEdit_3 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_3.setObjectName("lineEdit_3")
        self.verticalLayout.addWidget(self.lineEdit_3)
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setObjectName("pushButton_2")
        self.verticalLayout.addWidget(self.pushButton_2)
        self.pushButton_3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_3.setObjectName("pushButton_3")
        self.verticalLayout.addWidget(self.pushButton_3)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 313, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.pushButton_2.setText("Добавить")
        self.pushButton_3.setText("Отмена")
        self.pushButton_2.clicked.connect(self.setupadminisUi)

    def setupworkersUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(741, 380)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 0, 0, 1, 1)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 0, 1, 1, 4)
        self.pushButton_4 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_4.setObjectName("pushButton_4")
        self.gridLayout.addWidget(self.pushButton_4, 1, 6, 1, 1)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 1, 0, 1, 3)
        self.label_6 = QtWidgets.QLabel(self.centralwidget)
        self.label_6.setObjectName("label_6")
        self.gridLayout.addWidget(self.label_6, 2, 3, 1, 1)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 2, 0, 1, 2)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 2, 2, 1, 1)
        self.pushButton_6 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_6.setObjectName("pushButton_6")
        self.gridLayout.addWidget(self.pushButton_6, 1, 7, 1, 1)
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setObjectName("label_7")
        self.gridLayout.addWidget(self.label_7, 2, 5, 1, 1)
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 721, 253))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.gridLayout.addWidget(self.scrollArea, 3, 0, 1, 8)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 741, 20))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateworkersUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateworkersUi(self, adminis):
        _translate = QtCore.QCoreApplication.translate
        adminis.setWindowTitle(_translate("adminis", "MainWindow"))
        self.label.setText(_translate("adminis", "Вы авторизованы как:"))
        self.label_2.setText(_translate("adminis", "администратор"))
        self.pushButton_4.setText(_translate("adminis", "Закрыть"))
        self.label_3.setText(_translate("adminis", "<html><head/><body><p><span style=\" font-size:12pt;\">Список заказов:</span></p></body></html>"))
        self.label_6.setText(_translate("adminis", "Адрес"))
        self.label_4.setText(_translate("adminis", "Паспортные данные"))
        self.label_5.setText(_translate("adminis", "ФИО"))
        self.pushButton_6.setText(_translate("adminis", "Добавить "))
        self.label_7.setText(_translate("adminis", "Телефон"))

        global db_login

        cnx = mysql.connector.connect(user='root', password='i130813',
                                      host='127.0.0.1',
                                      database='aiskom')
        cursor = cnx.cursor()

        query = "select id_prodav, FIO, address, phone_numb from prodav;"
        cursor.execute(query)

        j = 0
        k = 0
        for item in cursor:
            if j == 0:
                j += 1
                continue
            for value in item:
                if k == 0:
                    line_item = QtWidgets.QLabel(str(value))
                    id = str(value)
                    self.gridLayout_2.addWidget(line_item, j, k, 1, 1)
                    k += 1
                    global special_id
                    special_id = id
                    continue
                line_item = QtWidgets.QLabel(str(value))
                self.gridLayout_2.addWidget(line_item, j, k, 1, 1)

                k += 1
                if k % 4 == 0:
                    but_item = QtWidgets.QPushButton("Открыть")
                    self.gridLayout_2.addWidget(but_item, j, 5, 1, 1)
                    # but_item.clicked.connect(lambda state, row=id: open_check(row))
                    but_item = QtWidgets.QPushButton("Удалить")
                    self.gridLayout_2.addWidget(but_item, j, 6, 1, 1)
                    # but_item.clicked.connect(lambda state, row=id: delete_dogovor(row))
                    j += 1
                    k = 0

        self.pushButton_6.clicked.connect(self.setupnewuserUi)

    def setupnewuserUi(self):
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.fioedit = QtWidgets.QLineEdit(self.centralwidget)
        self.fioedit.setObjectName("fioedit")
        self.verticalLayout.addWidget(self.fioedit)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.verticalLayout.addWidget(self.lineEdit_2)
        self.lineEdit_3 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_3.setObjectName("lineEdit_3")
        self.verticalLayout.addWidget(self.lineEdit_3)
        self.lineEdit_4 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_4.setObjectName("lineEdit_4")
        self.verticalLayout.addWidget(self.lineEdit_4)
        self.lineEdit_5 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_5.setObjectName("lineEdit_5")
        self.verticalLayout.addWidget(self.lineEdit_5)
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setObjectName("pushButton_2")
        self.verticalLayout.addWidget(self.pushButton_2)
        self.pushButton_3 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_3.setObjectName("pushButton_3")
        self.verticalLayout.addWidget(self.pushButton_3)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 313, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslatenewuserUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslatenewuserUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Dialog"))
        self.pushButton_2.setText(_translate("Dialog", "Добавить"))
        self.pushButton_3.setText(_translate("Dialog", "Назад"))

        self.pushButton_2.clicked.connect(lambda: adduser())

        def adduser():
            global db_login, db_host, db_pass, special_id

            cnx = mysql.connector.connect(user='root', password=db_pass,
                                          host=db_host,
                                          database='aiskom')
            cursor = cnx.cursor()
            if self.fioedit.text() != "" and self.lineEdit_2.text() != "" and self.lineEdit_3.text() != "" and self.lineEdit_4.text() != "" and self.lineEdit_5.text() != "":
                data = (self.fioedit.text(), self.lineEdit_2.text(), self.lineEdit_3.text(), self.lineEdit_4.text(), self.lineEdit_5.text())
                query = "insert into prodav(FIO, id_torgtoch, numb_ser_pass, phone_numb, address)values(%s,%s,%s,%s,%s);"
                cursor.execute(query, data)
                cnx.commit()
                for x in range(6):  # Количество символов (16)
                    pas = pas + random.choice(list('1234567890abcdefghigklmnopqrstuvyxwzABCDEFGHIGKLMNOPQRSTUVYXWZ'))
                data = (special_id, pas, special_id)
                query = "insert into uspas(login, password, id_prodav) values(%s,%s,%s);"
                cursor.execute(query, data)
                cnx.commit()
                self.setupworkersUi()

    def openotchet(self):
        Authorization = QtWidgets.QDialog()
        ui = Ui_MainWindow()
        ui.setupUi(Authorization)
        Authorization.exec_()

    def setupadministartorUi(self, administrator):
        administrator.setObjectName("administrator")
        administrator.resize(260, 211)
        self.centralwidget = QtWidgets.QWidget(administrator)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        administrator.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(administrator)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 260, 20))
        self.menubar.setObjectName("menubar")
        administrator.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(administrator)
        self.statusbar.setObjectName("statusbar")
        administrator.setStatusBar(self.statusbar)

        self.retranslateadministratorUi(administrator)
        QtCore.QMetaObject.connectSlotsByName(administrator)

    def retranslateadministratorUi(self, administrator):
        _translate = QtCore.QCoreApplication.translate
        administrator.setWindowTitle(_translate("administrator", "Dialog"))
        self.label.setText(_translate("administrator", "Отчет успешно загружен"))
        self.pushButton.setText(_translate("administrator", "Ок"))

    def setupUi(self, Authorization):
        Authorization.setObjectName("Authorization")
        Authorization.resize(282, 219)
        self.gridLayout = QtWidgets.QGridLayout(Authorization)
        self.gridLayout.setObjectName("gridLayout")
        self.label_3 = QtWidgets.QLabel(Authorization)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 0, 1, 1, 2)
        self.label_4 = QtWidgets.QLabel(Authorization)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 1, 1, 1, 2)
        self.label = QtWidgets.QLabel(Authorization)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 2, 0, 1, 1)
        self.lineEdit_2 = QtWidgets.QLineEdit(Authorization)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.gridLayout.addWidget(self.lineEdit_2, 2, 1, 1, 2)
        self.label_2 = QtWidgets.QLabel(Authorization)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 3, 0, 1, 1)
        self.lineEdit = QtWidgets.QLineEdit(Authorization)
        self.lineEdit.setObjectName("lineEdit")
        self.gridLayout.addWidget(self.lineEdit, 3, 1, 1, 2)
        self.label_5 = QtWidgets.QLabel(Authorization)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 4, 0, 1, 1)
        self.lineEdit_3 = QtWidgets.QLineEdit(Authorization)
        self.lineEdit_3.setObjectName("lineEdit_3")
        self.gridLayout.addWidget(self.lineEdit_3, 4, 1, 1, 2)
        self.pushButton = QtWidgets.QPushButton(Authorization)
        self.pushButton.setObjectName("pushButton")
        self.gridLayout.addWidget(self.pushButton, 6, 1, 1, 1)

        self.retranslateUi(Authorization)
        QtCore.QMetaObject.connectSlotsByName(Authorization)

    def retranslateUi(self, Authorization):
        _translate = QtCore.QCoreApplication.translate
        Authorization.setWindowTitle(_translate("Authorization", "Отчет"))
        self.label_3.setText(_translate("Authorization", "Формирование отчета"))
        self.label_4.setText(_translate("Authorization", "Статистика за период"))
        self.label.setText(_translate("Authorization", "С:"))
        self.lineEdit_2.setText(_translate("Authorization", "1-12-2018"))
        self.label_2.setText(_translate("Authorization", "По:"))
        self.lineEdit.setText(_translate("Authorization", "1-2-2019"))
        self.label_5.setText(_translate("Authorization", "Магазин:"))
        self.lineEdit_3.setText(_translate("Authorization", "Контрольная точка №2"))
        self.pushButton.setText(_translate("Authorization", "Продавцы"))

        self.pushButton.clicked.connect(self.otchettovary)

    def otchettovary(self):

        document = Document()
        document.add_heading('Отчет по продавцу Клавдин Сергей Анатольевич', 0)
        p = document.add_paragraph("Число продаж: 3")
        p = document.add_paragraph("На сумму: 48000")
        document.add_page_break()
        document.save('Отчет по продавцу.docx')





    def setupLoginUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(449, 315)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.verticalLayout.addWidget(self.label_3)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.verticalLayout.addWidget(self.label_4)
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setObjectName("lineEdit")
        self.verticalLayout.addWidget(self.lineEdit)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.verticalLayout.addWidget(self.label_2)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.verticalLayout.addWidget(self.label_5)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setMouseTracking(False)
        self.lineEdit_2.setTabletTracking(False)
        self.lineEdit_2.setContextMenuPolicy(QtCore.Qt.DefaultContextMenu)
        self.lineEdit_2.setToolTip("")
        self.lineEdit_2.setAutoFillBackground(False)
        self.lineEdit_2.setInputMethodHints(
            QtCore.Qt.ImhHiddenText | QtCore.Qt.ImhNoAutoUppercase | QtCore.Qt.ImhNoPredictiveText | QtCore.Qt.ImhSensitiveData)
        self.lineEdit_2.setEchoMode(QtWidgets.QLineEdit.Password)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.verticalLayout.addWidget(self.lineEdit_2)
        self.label_8 = QtWidgets.QLabel(self.centralwidget)
        self.label_8.setObjectName("label_8")
        self.verticalLayout.addWidget(self.label_8)
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setObjectName("label_7")
        self.verticalLayout.addWidget(self.label_7)
        self.lineEdit_3 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_3.setObjectName("lineEdit_3")
        self.verticalLayout.addWidget(self.lineEdit_3)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        self.label_6 = QtWidgets.QLabel(self.centralwidget)
        self.label_6.setObjectName("label_6")
        self.verticalLayout.addWidget(self.label_6)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 449, 21))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateLoginUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateLoginUi(self, oshibka):
        _translate = QtCore.QCoreApplication.translate
        oshibka.setWindowTitle(_translate("oshibka", "Авторизация"))
        self.label_3.setText(_translate("oshibka", "АИС для БД ООО \"Компьютер\""))
        self.label.setText(_translate("oshibka", "Логин:"))
        self.label_2.setText(_translate("oshibka", "Пароль:"))
        self.label_7.setText(_translate("oshibka", "Адрес:"))
        self.pushButton.setText(_translate("oshibka", "Войти"))
        self.label_6.setText(_translate("oshibka", ""))

        self.pushButton.clicked.connect(self.login)

    def login(self):

        global db_host, db_password, db_login

        if self.lineEdit.text() != "" and self.lineEdit_2.text() != "" and self.lineEdit_3.text() != "" and self.lineEdit_3.text().find(":") != -1:
            adress = self.lineEdit_3.text().split(":")
            password = str(adress[1])
            adress = str(adress[0])

            try:

                cnx = mysql.connector.connect(user='root', password=password,
                                          host=adress,
                                          database='aiskom')
                cursor = cnx.cursor()

                db_host = adress
                db_password = password
                db_login = self.lineEdit.text()

            except BaseException:
                self.label_6.setText("<html><head/><body><p><span style=\" color:#ff0000;\">Проверьте правильность введеных данных!!</span></p></body></html>")

            try:
                query = "select password from uspas where login= %s"
                data = (self.lineEdit.text(), )
                cursor.execute(query, data)
                for item in cursor:
                    if item[0] == self.lineEdit_2.text():
                        if self.lineEdit.text() == "admin":
                            self.setupadminisUi()
                        else:
                            self.setuporderUi()

            except BaseException:
               self.label_6.setText("<html><head/><body><p><span style=\" color:#ff0000;\">Проверьте правильность введеных данных!</span></p></body></html>")

        else:
            self.label_6.setText(
                "<html><head/><body><p><span style=\" color:#ff0000;\">Проверьте правильность введеных данных</span></p></body></html>")

    def setupsostavdogUi(self, sostavdog):
        sostavdog.setObjectName("sostavdog")
        sostavdog.resize(333, 417)
        self.gridLayout = QtWidgets.QGridLayout(sostavdog)
        self.gridLayout.setObjectName("gridLayout")
        self.pushButton_5 = QtWidgets.QPushButton(sostavdog)
        self.pushButton_5.setObjectName("pushButton_5")
        self.gridLayout.addWidget(self.pushButton_5, 0, 1, 1, 2)
        self.pushButton_4 = QtWidgets.QPushButton(sostavdog)
        self.pushButton_4.setObjectName("pushButton_4")
        self.gridLayout.addWidget(self.pushButton_4, 0, 3, 1, 1)
        self.label = QtWidgets.QLabel(sostavdog)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 1, 0, 1, 4)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.textEdit = QtWidgets.QTextEdit(sostavdog)
        self.textEdit.setObjectName("textEdit")
        self.horizontalLayout.addWidget(self.textEdit)
        self.textEdit_2 = QtWidgets.QTextEdit(sostavdog)
        self.textEdit_2.setObjectName("textEdit_2")
        self.horizontalLayout.addWidget(self.textEdit_2)
        self.textEdit_3 = QtWidgets.QTextEdit(sostavdog)
        self.textEdit_3.setObjectName("textEdit_3")
        self.horizontalLayout.addWidget(self.textEdit_3)
        self.gridLayout.addLayout(self.horizontalLayout, 2, 0, 1, 4)
        self.pushButton = QtWidgets.QPushButton(sostavdog)
        self.pushButton.setObjectName("pushButton")
        self.gridLayout.addWidget(self.pushButton, 3, 0, 1, 1)
        self.pushButton_2 = QtWidgets.QPushButton(sostavdog)
        self.pushButton_2.setObjectName("pushButton_2")
        self.gridLayout.addWidget(self.pushButton_2, 3, 1, 1, 1)
        self.pushButton_3 = QtWidgets.QPushButton(sostavdog)
        self.pushButton_3.setObjectName("pushButton_3")
        self.gridLayout.addWidget(self.pushButton_3, 3, 2, 1, 2)

        self.retranslatesostavdogUi(sostavdog)
        QtCore.QMetaObject.connectSlotsByName(sostavdog)

    def retranslatesostavdogUi(self, sostavdog):
        _translate = QtCore.QCoreApplication.translate
        sostavdog.setWindowTitle(_translate("sostavdog", "Dialog"))
        self.pushButton_5.setText(_translate("sostavdog", "Добавить"))
        self.pushButton_4.setText(_translate("sostavdog", "Закрыть"))
        self.label.setText(_translate("sostavdog", "Штрафы и премии (дата, объем)"))
        self.textEdit.setHtml(_translate("sostavdog", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:8.25pt; font-weight:400; font-style:normal;\">\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">12.04.2017</p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">-1299</p></body></html>"))
        self.textEdit_2.setHtml(_translate("sostavdog", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:8.25pt; font-weight:400; font-style:normal;\">\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">13.05.2017</p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">23000</p></body></html>"))
        self.textEdit_3.setHtml(_translate("sostavdog", "<!DOCTYPE HTML PUBLIC \"-//W3C//DTD HTML 4.0//EN\" \"http://www.w3.org/TR/REC-html40/strict.dtd\">\n"
"<html><head><meta name=\"qrichtext\" content=\"1\" /><style type=\"text/css\">\n"
"p, li { white-space: pre-wrap; }\n"
"</style></head><body style=\" font-family:\'MS Shell Dlg 2\'; font-size:8.25pt; font-weight:400; font-style:normal;\">\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">24.09.2017</p>\n"
"<p style=\" margin-top:0px; margin-bottom:0px; margin-left:0px; margin-right:0px; -qt-block-indent:0; text-indent:0px;\">12000</p></body></html>"))
        self.pushButton.setText(_translate("sostavdog", "Убрать"))
        self.pushButton_2.setText(_translate("sostavdog", "Убрать"))
        self.pushButton_3.setText(_translate("sostavdog", "Убрать"))

    def setuporderUi(self):
        MainWindow.setObjectName("adminis")
        MainWindow.resize(741, 380)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.workerlabel = QtWidgets.QLabel(self.centralwidget)
        self.workerlabel.setObjectName("workerlabel")
        self.gridLayout.addWidget(self.workerlabel, 0, 0, 1, 1)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 0, 1, 1, 4)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 1, 0, 1, 3)
        self.pushButton_4 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_4.setObjectName("pushButton_4")
        self.gridLayout.addWidget(self.pushButton_4, 1, 6, 1, 1)
        self.labelorderdate = QtWidgets.QLabel(self.centralwidget)
        self.labelorderdate.setObjectName("labelorderdate")
        self.gridLayout.addWidget(self.labelorderdate, 2, 3, 1, 1)
        self.pushButton_6 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_6.setObjectName("pushButton_6")
        self.gridLayout.addWidget(self.pushButton_6, 1, 7, 1, 1)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 2, 2, 1, 1)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 2, 0, 1, 2)
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setObjectName("label_7")
        self.gridLayout.addWidget(self.label_7, 2, 5, 1, 1)
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 721, 253))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.gridLayout.addWidget(self.scrollArea, 3, 0, 1, 8)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 741, 20))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateorderUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateorderUi(self, orderui):
        _translate = QtCore.QCoreApplication.translate
        orderui.setWindowTitle(_translate("orderui", "MainWindow"))
        self.workerlabel.setText(_translate("orderui", "Вы авторизованы как:"))
        #self.label_2.setText(_translate("orderui", "администратор"))
        self.label_3.setText(_translate("orderui", "<html><head/><body><p><span style=\" font-size:12pt;\">Список заказов:</span></p></body></html>"))
        self.pushButton_4.setText(_translate("orderui", "Закрыть"))
        self.labelorderdate.setText(_translate("orderui", "Дата заказа"))
        self.pushButton_6.setText(_translate("orderui", "Добавить "))
        self.label_5.setText(_translate("orderui", "Кол-во товара"))
        self.label_4.setText(_translate("orderui", "Номер заказа"))
        self.label_7.setText(_translate("orderui", "Стоимость"))

        self.pushButton_6.clicked.connect(lambda: add())

        def add():
            first_column_ord.append('120003')
            second_column_ord.append('24-01-2019')
            third_column_ord.append('1')
            four_column_ord.append('30000')
            self.setupaddorderUi()

        i = 0
        for item in first_column_ord:
            line_item = QtWidgets.QLabel(str(item))
            self.gridLayout_2.addWidget(line_item, i, 1, 1, 1)
            i += 1
        i = 0
        for item in second_column_ord:
            line_item = QtWidgets.QLineEdit(str(item))
            self.gridLayout_2.addWidget(line_item, i, 2, 1, 1)
            i += 1
        i = 0
        for item in third_column_ord:
            line_item = QtWidgets.QLineEdit(str(item))
            self.gridLayout_2.addWidget(line_item, i, 3, 1, 1)
            line_item = QtWidgets.QPushButton("Удалить")
            line_item.clicked.connect(lambda state, row = i: delete(row))
            self.gridLayout_2.addWidget(line_item, i, 5, 1, 1)
            i += 1
        i = 0
        for item in four_column_ord:
            line_item = QtWidgets.QLineEdit(str(item))
            self.gridLayout_2.addWidget(line_item, i, 4, 1, 1)
            i += 1

        def delete(row):
            first_column_ord.pop(row)
            second_column_ord.pop(row)
            third_column_ord.pop(row)
            four_column_ord.pop(row)
            self.setuporderUi()

    def setupaddorderUi(self):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(260, 211)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setObjectName("lineEdit")
        self.verticalLayout.addWidget(self.lineEdit)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.verticalLayout.addWidget(self.lineEdit_2)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setObjectName("pushButton_2")
        self.verticalLayout.addWidget(self.pushButton_2)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 260, 20))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.pushButton.setText("Добавить")
        self.pushButton_2.setText("Отмена")
        self.pushButton.clicked.connect(self.setuporderUi)

    def setupcheckUi(self, checkui):
        checkui.setObjectName("adminis")
        checkui.resize(741, 380)
        self.centralwidget = QtWidgets.QWidget(checkui)
        self.centralwidget.setObjectName("centralwidget")
        self.gridLayout = QtWidgets.QGridLayout(self.centralwidget)
        self.gridLayout.setObjectName("gridLayout")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.gridLayout.addWidget(self.label, 0, 0, 1, 1)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setObjectName("label_2")
        self.gridLayout.addWidget(self.label_2, 0, 1, 1, 4)
        self.pushButton_9 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_9.setObjectName("pushButton_9")
        self.gridLayout.addWidget(self.pushButton_9, 1, 3, 1, 3)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.gridLayout.addWidget(self.label_3, 1, 0, 1, 3)
        self.pushButton_4 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_4.setObjectName("pushButton_4")
        self.gridLayout.addWidget(self.pushButton_4, 1, 6, 1, 1)
        self.pushButton_6 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_6.setObjectName("pushButton_6")
        self.gridLayout.addWidget(self.pushButton_6, 1, 7, 1, 1)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.gridLayout.addWidget(self.label_5, 2, 2, 1, 1)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.gridLayout.addWidget(self.label_4, 2, 0, 1, 2)
        self.scrollArea = QtWidgets.QScrollArea(self.centralwidget)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 721, 253))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.gridLayout_2 = QtWidgets.QGridLayout(self.scrollAreaWidgetContents)
        self.gridLayout_2.setObjectName("gridLayout_2")
        self.lineEdit = QtWidgets.QLineEdit(self.scrollAreaWidgetContents)
        self.lineEdit.setObjectName("lineEdit")
        self.gridLayout_2.addWidget(self.lineEdit, 0, 0, 1, 1)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.scrollAreaWidgetContents)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.gridLayout_2.addWidget(self.lineEdit_2, 0, 1, 1, 1)
        self.pushButton = QtWidgets.QPushButton(self.scrollAreaWidgetContents)
        self.pushButton.setObjectName("pushButton")
        self.gridLayout_2.addWidget(self.pushButton, 0, 4, 1, 1)
        self.pushButton_2 = QtWidgets.QPushButton(self.scrollAreaWidgetContents)
        self.pushButton_2.setObjectName("pushButton_2")
        self.gridLayout_2.addWidget(self.pushButton_2, 0, 5, 1, 1)
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.gridLayout.addWidget(self.scrollArea, 3, 0, 1, 8)
        checkui.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(checkui)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 741, 20))
        self.menubar.setObjectName("menubar")
        checkui.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(checkui)
        self.statusbar.setObjectName("statusbar")
        checkui.setStatusBar(self.statusbar)

        self.retranslatecheckUi(checkui)
        QtCore.QMetaObject.connectSlotsByName(checkui)

    def retranslatecheckUi(self, checkui):
        _translate = QtCore.QCoreApplication.translate
        checkui.setWindowTitle(_translate("checkui", "MainWindow"))
        self.label.setText(_translate("checkui", ""))

        global db_login

        cnx = mysql.connector.connect(user='root', password='i130813',
                                      host='127.0.0.1',
                                      database='aiskom')
        cursor = cnx.cursor()

        query = "select FIO from prodav where id_prodav=%s"
        data = (db_login, )
        cursor.execute(query, data)

        for item in query:
            for value in item:
                self.label.setText("Вы авторизованы как: " + str(value))

        query = "select name from torgtoch,prodav where id_prodav=%s and torgtoch.id_torgtoch=prodav.id_torgtoch;"
        data = (db_login, )
        cursor.execute(query, data)
        for item in query:
            for value in item:
                self.label.setText("Магазин: " + str(value))

        self.pushButton_9.setText(_translate("checkui", "Сформировать отчет"))
        self.label_3.setText(_translate("checkui", "<html><head/><body><p><span style=\" font-size:12pt;\">Список чеков:</span></p></body></html>"))
        self.pushButton_4.setText(_translate("checkui", "Закрыть"))
        self.pushButton_6.setText(_translate("checkui", "Добавить "))
        self.label_5.setText(_translate("checkui", "Дата продажи"))
        self.label_4.setText(_translate("checkui", "Номер продажи"))
        self.pushButton.setText(_translate("checkui", "Открыть"))
        self.pushButton_2.setText(_translate("checkui", "Удалить"))

        self.pushButton_6.clicked.connect(lambda: add())

        def add():
             first_column_chek.append('110003')
             second_column_chek.append('24-01-2019')
             self.setupaddcheckUi()

        i = 0
        for item in first_column_chek:
            line_item = QtWidgets.QLabel(str(item))
            self.gridLayout_2.addWidget(line_item, i, 1, 1, 1)
            i += 1
        i = 0
        for item in second_column_chek:
            line_item = QtWidgets.QLineEdit(str(item))
            self.gridLayout_2.addWidget(line_item, i, 2, 1, 1)
            line_item = QtWidgets.QPushButton("Удалить")
            line_item.clicked.connect(lambda state, row = i: delete(row))
            self.gridLayout_2.addWidget(line_item, i, 3, 1, 1)
            i += 1

        def delete(row):
            first_column_chek.pop(row)
            second_column_chek.pop(row)
            self.setupcheckUi()

    def setupaddcheckUi(self, administrator):
        administrator.setObjectName("administrator")
        administrator.resize(260, 211)
        self.centralwidget = QtWidgets.QWidget(administrator)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.lineEdit = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit.setObjectName("lineEdit")
        self.verticalLayout.addWidget(self.lineEdit)
        self.lineEdit_2 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_2.setObjectName("lineEdit_2")
        self.verticalLayout.addWidget(self.lineEdit_2)
        self.lineEdit_3 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_3.setObjectName("lineEdit_3")
        self.verticalLayout.addWidget(self.lineEdit_3)
        self.lineEdit_4 = QtWidgets.QLineEdit(self.centralwidget)
        self.lineEdit_4.setObjectName("lineEdit_4")
        self.verticalLayout.addWidget(self.lineEdit_4)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setObjectName("pushButton_2")
        self.verticalLayout.addWidget(self.pushButton_2)
        administrator.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(administrator)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 260, 20))
        self.menubar.setObjectName("menubar")
        administrator.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(administrator)
        self.statusbar.setObjectName("statusbar")
        administrator.setStatusBar(self.statusbar)

        self.retranslateaddcheckUi(administrator)
        QtCore.QMetaObject.connectSlotsByName(administrator)

    def retranslateaddcheckUi(self, administrator):
        _translate = QtCore.QCoreApplication.translate
        administrator.setWindowTitle(_translate("administrator", "Dialog"))
        self.pushButton.setText(_translate("administrator", "Ок"))
        self.pushButton_2.setText(_translate("administrator", "Отмена"))

        self.pushButton.clicked.connect(self.setupcheckUi)

    def setupcheckviewUi(self, administrator):
        administrator.setObjectName("administrator")
        administrator.resize(260, 211)
        self.centralwidget = QtWidgets.QWidget(administrator)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout.setObjectName("verticalLayout")
        self.workerlabel = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.pushButton = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton.setObjectName("pushButton")
        self.verticalLayout.addWidget(self.pushButton)
        administrator.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(administrator)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 260, 20))
        self.menubar.setObjectName("menubar")
        administrator.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(administrator)
        self.statusbar.setObjectName("statusbar")
        administrator.setStatusBar(self.statusbar)

        self.retranslatecheckviewUi(administrator)
        QtCore.QMetaObject.connectSlotsByName(administrator)

    def retranslatecheckviewUi(self, administrator):
        _translate = QtCore.QCoreApplication.translate
        administrator.setWindowTitle(_translate("administrator", "Dialog"))
        self.pushButton.setText(_translate("administrator", "Ок"))

        global db_login

        cnx = mysql.connector.connect(user='root', password='i130813',
                                      host='127.0.0.1',
                                      database='aiskom')
        cursor = cnx.cursor()

        query = "select * from prodazha;"

        cursor.execute(query)

        k = 0
        for item in query:
            for value in item:
                if k == 0:
                    line_item = QtWidgets.QLabel(str(value))
                    self.scrollAreaWidgetContents.addWidget(line_item)
                    k += 1
                    continue
                line_item = QtWidgets.QLineEdit(str(value))
                self.scrollAreaWidgetContents.addWidget(line_item)


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MainWindow()
    ui.setupadminisUi()
    MainWindow.show()
    sys.exit(app.exec_())

